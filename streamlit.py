import os
import math
import time
import logging
import ffmpeg
import tempfile
import streamlit as st
from dotenv import load_dotenv
from pydub import AudioSegment
from sarvamai import SarvamAI
from docx import Document
from fpdf import FPDF
from langchain_google_genai import ChatGoogleGenerativeAI

# Setup logging
logging.basicConfig(level=logging.INFO)

# Load environment variables
load_dotenv()
SARVAM_KEY = os.getenv("SARVAM_SUBSCRIPTION_KEY")
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")

client = SarvamAI(api_subscription_key=SARVAM_KEY)

def convert_to_wav(input_path, output_path="converted_audio.wav"):
    ffmpeg.input(input_path).output(output_path, ac=1, ar=16000).overwrite_output().run()
    return output_path

def split_audio(file_path, chunk_length_sec=30):
    audio = AudioSegment.from_wav(file_path)
    chunks = []
    base = os.path.splitext(file_path)[0]
    total_chunks = math.ceil(len(audio) / (chunk_length_sec * 1000))

    for i in range(total_chunks):
        start = i * chunk_length_sec * 1000
        end = min((i + 1) * chunk_length_sec * 1000, len(audio))
        chunk = audio[start:end]
        chunk_path = f"{base}_part{i+1}.wav"
        chunk.export(chunk_path, format="wav")
        chunks.append(chunk_path)

    return chunks

def transcribe_audio(audio_path):
    chunks = split_audio(audio_path)
    transcript = []

    for chunk in chunks:
        with open(chunk, "rb") as f:
            try:
                result = client.speech_to_text.translate(file=f, model="saaras:v2")
                transcript.append(result.transcript)
            except Exception as e:
                st.error(f"Error transcribing chunk {chunk}: {e}")
        os.remove(chunk)
        time.sleep(0.5) 

    return " ".join(transcript)

def summarize_with_gemini(text):
    chunk_size = 30000
    chunks = [text[i:i + chunk_size] for i in range(0, len(text), chunk_size)]
    all_summaries = []

    for chunk in chunks:
        llm = ChatGoogleGenerativeAI(
            model="gemini-2.0-flash",
            temperature=0.7,
            max_tokens=2048,
            google_api_key=GEMINI_API_KEY
        )
        prompt = (
            "Summarize the following meeting transcript into bullet points grouped by speaker names. "
            "Do not include any introductions or headings.\n\n"
            f"{chunk}"
        )
        try:
            summary = llm.invoke(prompt)
            summary_text = summary.content if hasattr(summary, 'content') else str(summary)
            all_summaries.append(summary_text)
        except Exception as e:
            all_summaries.append(f"[Error summarizing chunk]: {e}")

    return "\n\n".join(all_summaries)

def save_docx(text, title="Document"):
    doc = Document()
    doc.add_heading(title, 0)
    doc.add_paragraph(text)
    tmp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(tmp_file.name)
    return tmp_file.name

def save_pdf(text, title="Document"):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", "B", 16)
    pdf.multi_cell(0, 10, title)
    pdf.set_font("Arial", "", 12)
    for paragraph in text.split("\n"):
        pdf.multi_cell(0, 10, paragraph)
    tmp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
    pdf.output(tmp_file.name)
    return tmp_file.name

# Streamlit UI
st.title("Audio/Video MultiLingual Summarizer")

uploaded_file = st.file_uploader("Upload your audio or video file", type=["mp3", "wav", "m4a", "aac", "mp4", "mov", "mkv"])
if uploaded_file is not None:

    with tempfile.NamedTemporaryFile(delete=False, suffix=uploaded_file.name[-4:]) as temp_input:
        temp_input.write(uploaded_file.read())
        temp_input.flush()
        wav_path = convert_to_wav(temp_input.name)

    st.info("Converting and transcribing...")

    start_transcribe = time.time()
    transcript = transcribe_audio(wav_path)
    end_transcribe = time.time()
    st.success(f"Transcription completed in {round(end_transcribe - start_transcribe, 2)} seconds.")

    st.subheader("Transcript")
    st.text_area("Full Transcript", transcript, height=300)

    transcript_docx_path = save_docx(transcript, title="Transcript")
    transcript_pdf_path = save_pdf(transcript, title="Transcript")
    left, col1, col2, right = st.columns([1, 2, 2, 1])
    with col1:
        st.download_button("Transcript DOCX", open(transcript_docx_path, "rb"), file_name="transcript.docx")
    with col2:
        st.download_button("Transcript PDF", open(transcript_pdf_path, "rb"), file_name="transcript.pdf")

    os.remove("converted_audio.wav")

    st.info("Summarizing with Gemini...")

    start_summary = time.time()
    summary = summarize_with_gemini(transcript)
    end_summary = time.time()
    st.success(f"Summary generated in {round(end_summary - start_summary, 2)} seconds.")

    st.subheader("Summary")
    st.text_area("Meeting Summary", summary, height=250)

    summary_docx_path = save_docx(summary, title="Summary")
    summary_pdf_path = save_pdf(summary, title="Summary")
    left, col3, col4, right = st.columns([1, 2, 2, 1])
    with col3:
        st.download_button("Summary DOCX", open(summary_docx_path, "rb"), file_name="summary.docx")
    with col4:
        st.download_button("Summary PDF", open(summary_pdf_path, "rb"), file_name="summary.pdf")
